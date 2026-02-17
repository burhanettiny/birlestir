import os
import tempfile
from io import BytesIO
import streamlit as st
from docx import Document
from pypdf import PdfMerger, PdfReader, PdfWriter

# ---------------------------
# Session state başlangıcı
# ---------------------------
# processed_pdfs: Düzenlenmiş (sayfası silinmiş) PDF'lerin bayt verisini tutar
if "processed_pdfs" not in st.session_state:
    st.session_state.processed_pdfs = {}

# uploaded_meta: Yüklenen dosyaların bilgilerini tutar
if "uploaded_meta" not in st.session_state:
    st.session_state.uploaded_meta = []

# ---------------------------
# Streamlit UI Ayarları
# ---------------------------
st.set_page_config(page_title="Belge Birleştirici", page_icon="📎", layout="centered")
st.title("📎 PDF & Word Birleştirici")

# --- YAN PANEL (SIDEBAR) ---
st.sidebar.header("⚙️ Kontrol Paneli")
if st.sidebar.button("🗑️ Tüm Listeyi Sıfırla"):
    st.session_state.uploaded_meta = []
    st.session_state.processed_pdfs = {}
    st.rerun()

st.sidebar.markdown("---")
st.sidebar.info("""
**Nasıl Çalışır?**
1. Dosyaları yükleyin.
2. Sıralamayı multiselect ile belirleyin.
3. PDF'lerde sayfa silme işlemi yapın (Opsiyonel).
4. Birleştir ve İndir butonuna basın.
""")

# --- DOSYA YÜKLEME ---
uploaded_files = st.file_uploader(
    "PDF veya Word dosyalarını yükleyin (Çoklu seçim yapabilirsiniz)",
    type=["pdf", "docx"],
    accept_multiple_files=True
)

# ---------------------------
# Dosya Senkronizasyon Mantığı
# ---------------------------
if uploaded_files:
    # Uploader'daki dosyaların kimliklerini (isim_boyut) oluştur
    uploader_file_ids = [f"{f.name}_{f.size}" for f in uploaded_files]
    
    # 1. Yeni yüklenenleri ekle
    current_meta_keys = [m["key"] for m in st.session_state.uploaded_meta]
    for f in uploaded_files:
        f_key = f"{f.name}_{f.size}"
        if f_key not in current_meta_keys:
            st.session_state.uploaded_meta.append({
                "key": f_key,
                "name": f.name,
                "file": f
            })
    
    # 2. Uploader'dan kaldırılanları meta listesinden de sil
    # (Kullanıcı uploader kutusundaki 'x'e basarsa listeden gider)
    st.session_state.uploaded_meta = [
        m for m in st.session_state.uploaded_meta if m["key"] in uploader_file_ids
    ]

# Liste boşsa durdur
if not st.session_state.uploaded_meta:
    st.info("Lütfen işlem yapmak için dosya yükleyin.")
    st.stop()

# ---------------------------
# Sıralama Yönetimi
# ---------------------------
st.subheader("🗂️ 1. Sıralama ve Dosya Listesi")
# ID ekleyerek aynı isimli dosyaların karışmasını engelliyoruz
choices = [f'{m["name"]} (ID: {i})' for i, m in enumerate(st.session_state.uploaded_meta)]
sorted_choice = st.multiselect(
    "Birleştirme sırasını belirleyin (Üstten alta doğru birleşir):",
    choices,
    default=choices
)

# Seçilen sıraya göre listeyi yeniden oluştur
ordered_indices = [int(c.split("(ID: ")[-1].strip(")")) for c in sorted_choice]
sorted_meta = [st.session_state.uploaded_meta[i] for i in ordered_indices]

# ---------------------------
# PDF Sayfa Silme / Düzenleme
# ---------------------------
pdf_meta_list = [m for m in sorted_meta if m["name"].lower().endswith(".pdf")]

if pdf_meta_list:
    st.markdown("---")
    st.subheader("📄 2. PDF Sayfa Yönetimi (Opsiyonel)")
    pdf_to_edit_name = st.selectbox("Düzenlemek istediğiniz PDF'i seçin", ["Seçiniz"] + [m["name"] for m in pdf_meta_list])
    
    if pdf_to_edit_name != "Seçiniz":
        selected_meta = next(m for m in pdf_meta_list if m["name"] == pdf_to_edit_name)
        f_obj = selected_meta["file"]
        f_obj.seek(0)
        
        reader = PdfReader(f_obj)
        total_pages = len(reader.pages)
        st.write(f"💡 **{selected_meta['name']}** toplam {total_pages} sayfa.")
        
        delete_pages = st.multiselect("Silinecek sayfaları seçin:", [f"Sayfa {i+1}" for i in range(total_pages)])

        if st.button("📌 Seçili Sayfaları Çıkar ve Kaydet"):
            writer = PdfWriter()
            for idx in range(total_pages):
                if f"Sayfa {idx+1}" not in delete_pages:
                    writer.add_page(reader.pages[idx])
            
            buf = BytesIO()
            writer.write(buf)
            st.session_state.processed_pdfs[selected_meta["key"]] = buf.getvalue()
            st.success(f"'{selected_meta['name']}' güncellendi. Birleştirmede bu hali kullanılacak.")

# ---------------------------
# BİRLEŞTİRME VE İNDİRME
# ---------------------------
st.markdown("---")
st.subheader("🚀 3. Birleştir ve İndir")

col1, col2 = st.columns(2)

with col1:
    pdf_to_merge = [m for m in sorted_meta if m["name"].lower().endswith(".pdf")]
    if st.button("🔀 PDF'leri Birleştir", disabled=not pdf_to_merge, use_container_width=True):
        try:
            merger = PdfMerger()
            for m in pdf_to_merge:
                # Eğer düzenlenmiş hali varsa onu, yoksa orijinali kullan
                data = st.session_state.processed_pdfs.get(m["key"], m["file"].getvalue())
                merger.append(BytesIO(data))
            
            out_pdf = BytesIO()
            merger.write(out_pdf)
            st.download_button("📥 Birleşmiş PDF'i İndir", out_pdf.getvalue(), "merged_result.pdf", "application/pdf")
        except Exception as e:
            st.error(f"Hata: {e}")

with col2:
    docx_to_merge = [m for m in sorted_meta if m["name"].lower().endswith(".docx")]
    if st.button("📝 Word'leri Birleştir", disabled=not docx_to_merge, use_container_width=True):
        try:
            merged_doc = Document()
            for i, m in enumerate(docx_to_merge):
                if i > 0: merged_doc.add_page_break()
                sub_doc = Document(BytesIO(m["file"].getvalue()))
                for p in sub_doc.paragraphs:
                    merged_doc.add_paragraph(p.text)
            
            out_docx = BytesIO()
            merged_doc.save(out_docx)
            st.download_button("📥 Birleşmiş Word İndir", out_docx.getvalue(), "merged_result.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        except Exception as e:
            st.error(f"Hata: {e}")

st.markdown("---")
st.caption("Verileriniz oturum bazlıdır, sayfayı yenilediğinizde veya kapattığınızda silinir.")
